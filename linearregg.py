import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches

data = pd.read_csv('ml_fitbit.csv')


t = data['TotalSteps'].values.reshape(-1, 1)
b = data['Calories'].values.reshape(-1, 1)
n = 4


for i in range(4, n):
    t = np.hstack((t, t[:, 0] ** i))

A = np.hstack((np.ones((len(t), 1)), t))
ATA = np.dot(A.T, A)
ATb = np.dot(A.T, b)

x = np.linalg.solve(ATA, ATb)


y_pred = np.dot(A, x)
y_mean = np.mean(b)
ss_tot = np.sum((b - y_mean) ** 2)
ss_res = np.sum((b - y_pred) ** 2)
mse = ss_res / len(b)
rmse = np.sqrt(mse)
mae = np.mean(np.abs(b - y_pred))
r_squared = 1 - (ss_res / ss_tot)


doc = Document()

doc.add_heading('Univariate Linear Regression Metrics', level=0)

doc.add_paragraph(f'Savele Costin')
doc.add_paragraph(f'R-squared: {r_squared:.4f}')
doc.add_paragraph(f'SSE: {ss_res:.4f}')
doc.add_paragraph(f'MSE: {mse:.4f}')
doc.add_paragraph(f'RMSE: {rmse:.4f}')
doc.add_paragraph(f'MAE: {mae:.4f}')

doc.save('univariate_regression_metrics.docx')
